import streamlit as st
from openai import OpenAI
import json
import urllib.parse
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

# 页面配置
st.set_page_config(
    page_title="教学设计智能体",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded"
)

# 自定义CSS样式
st.markdown("""
<style>
    /* 主题色彩 */
    :root {
        --primary-color: #2E86AB;
        --secondary-color: #A23B72;
        --accent-color: #F18F01;
    }
    
    /* 标题样式 */
    .main-title {
        font-family: 'Segoe UI', sans-serif;
        font-size: 3rem;
        font-weight: 700;
        background: linear-gradient(120deg, #2E86AB, #A23B72, #F18F01);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        text-align: center;
        padding: 1.5rem 0;
        margin-bottom: 0.5rem;
    }
    
    .sub-title {
        text-align: center;
        color: #666;
        font-size: 1.2rem;
        margin-bottom: 2rem;
        font-weight: 300;
    }
    
    /* 卡片样式 */
    .info-card {
        background: linear-gradient(135deg, #667eea15 0%, #764ba215 100%);
        border-left: 4px solid #2E86AB;
        padding: 1.5rem;
        border-radius: 10px;
        margin: 1rem 0;
        box-shadow: 0 2px 8px rgba(0,0,0,0.1);
    }
    
    /* 按钮动画 */
    .stButton>button {
        transition: all 0.3s ease;
        border-radius: 8px;
        font-weight: 600;
    }
    
    .stButton>button:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(0,0,0,0.15);
    }
    
    /* 输入框样式 */
    .stTextInput>div>div>input, .stTextArea>div>div>textarea {
        border-radius: 8px;
        border: 2px solid #e0e0e0;
        transition: border-color 0.3s ease;
    }
    
    .stTextInput>div>div>input:focus, .stTextArea>div>div>textarea:focus {
        border-color: #2E86AB;
        box-shadow: 0 0 0 2px rgba(46, 134, 171, 0.1);
    }
    
    /* 成功/错误消息样式 */
    .stSuccess, .stError, .stWarning, .stInfo {
        border-radius: 8px;
        padding: 1rem;
        animation: slideIn 0.3s ease;
    }
    
    @keyframes slideIn {
        from {
            opacity: 0;
            transform: translateY(-10px);
        }
        to {
            opacity: 1;
            transform: translateY(0);
        }
    }
    
    /* 分隔线样式 */
    hr {
        margin: 2rem 0;
        border: none;
        height: 2px;
        background: linear-gradient(90deg, transparent, #2E86AB, transparent);
    }
    
    /* 标签样式 */
    .stSelectbox label, .stTextInput label, .stTextArea label, .stSlider label {
        font-weight: 600;
        color: #2E86AB;
    }
</style>
""", unsafe_allow_html=True)

# 教学方法定义
TEACHING_METHODS = {
    "讲授法": "教师系统讲解知识，学生听讲、记录和理解",
    "翻转课堂": "学生课前自学，课堂进行讨论、实践和答疑",
    "项目式学习": "以真实项目为驱动，学生通过完成项目掌握知识",
    "探究式学习": "引导学生提出问题、探索和发现知识",
    "合作学习": "学生分组协作，共同完成学习任务",
    "案例教学": "通过分析真实案例，培养学生解决问题的能力",
    "游戏化学习": "将游戏元素融入教学，提高学习兴趣和参与度"
}

def generate_teaching_design(client, subject, content, method, duration, audience, 
                            key_points, difficulties, teaching_goals, student_level, 
                            additional_requirements, custom_duration=None):
    """生成教学设计"""
    
    system_prompt = """你是一位资深的物理教学设计专家，擅长为不同学科和教学方法设计高质量的教学过程。
你需要根据用户提供的信息，设计详细的教学过程。

**重要要求**：
1. 教学目标必须基于物理核心素养（物理观念、科学思维、科学探究、科学态度与责任）进行设计，不能直接照搬用户输入的教学目标，而是要结合用户目标和核心素养进行优化
2. 必须包含详细的板书设计部分
3. 教学过程必须使用Markdown表格呈现，表格包含列：教学环节 | 时间 | 教师活动 | 学生活动 | 设计意图
4. 教学评价方式

请以结构化的方式输出，使用 Markdown 格式，教学过程部分必须使用Markdown表格。"""

    # 构建用户提示
    user_prompt = f"""请为以下教学内容设计教学过程：

**学科/课程**: {subject}
**授课内容**: {content}
**教学方法**: {method} ({TEACHING_METHODS[method]})
**课时时长**: {custom_duration if custom_duration else duration}分钟
**教学对象**: {audience}
**学生情况**: {student_level}

**教学重点**: {key_points}
**教学难点**: {difficulties}

**用户期望的教学目标**（需要结合物理核心素养优化）:
{teaching_goals}

**其他要求**:
{additional_requirements}

请设计详细的教学过程，确保：
1. 教学目标基于物理核心素养（物理观念、科学思维、科学探究、科学态度与责任）进行设计，结合用户期望但不照搬
2. 根据学生情况（{student_level}）因材施教
3. 包含完整的板书设计
4. 符合{method}的特点和要求
5. 满足用户的其他要求"""

    response = client.chat.completions.create(
        model="deepseek-chat",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ],
        stream=True,
        temperature=0.7
    )
    
    return response

def generate_ppt_outline(client, subject, content, method, teaching_design):
    """生成PPT大纲"""
    
    system_prompt = """你是一位专业的PPT设计师，擅长为教学内容设计清晰、美观的PPT演示文稿大纲。
你需要根据教学设计，创建PPT大纲，包括：
1. 每一页的标题
2. 每一页的主要内容要点
3. 建议的视觉元素（图片、图表、动画等）
4. 页面布局建议

请以结构化的方式输出，使用 Markdown 格式，教学过程部分必须使用Markdown表格。"""

    user_prompt = f"""基于以下教学设计，请创建PPT演示文稿大纲：

**学科/课程**: {subject}
**授课内容**: {content}
**教学方法**: {method}

**教学设计**:
{teaching_design}

请设计15-25页的PPT大纲，确保内容完整、逻辑清晰、视觉效果好。"""

    response = client.chat.completions.create(
        model="deepseek-chat",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ],
        stream=True,
        temperature=0.7
    )
    
    return response


def markdown_to_word(markdown_text, title):
    """将Markdown文本转换为Word文档，支持表格"""
    doc = Document()
    
    # 设置标题
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 先把文本按行分割，检测表格块
    lines = markdown_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        # 检测表格（以 | 开头）
        if line.startswith("|") and "|" in line[1:]:
            # 收集整个表格
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            
            # 过滤掉分隔行（如 |---|---|）
            data_lines = [l for l in table_lines if not re.match(r"^\|[-:\s|]+\|$", l)]
            
            if len(data_lines) >= 1:
                # 解析表头
                headers = [cell.strip() for cell in data_lines[0].split("|") if cell.strip()]
                n_cols = len(headers)
                n_rows = len(data_lines)
                
                # 创建 Word 表格
                table = doc.add_table(rows=n_rows, cols=n_cols)
                table.style = "Table Grid"
                
                # 填入数据
                for row_idx, row_line in enumerate(data_lines):
                    cells = [cell.strip() for cell in row_line.split("|") if cell.strip()]
                    for col_idx, cell_text in enumerate(cells[:n_cols]):
                        cell = table.cell(row_idx, col_idx)
                        cell.text = cell_text
                        if row_idx == 0:
                            # 表头加粗
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    run.bold = True
                
                doc.add_paragraph()  # 表格后空行
            continue
        
        # 处理标题
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=4)
        # 处理列表
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif re.match(r"^\d+\. ", line):
            doc.add_paragraph(re.sub(r"^\d+\. ", "", line), style="List Number")
        # 处理加粗
        elif "**" in line:
            p = doc.add_paragraph()
            parts = line.split("**")
            for idx, part in enumerate(parts):
                if idx % 2 == 0:
                    p.add_run(part)
                else:
                    p.add_run(part).bold = True
        # 普通段落
        else:
            doc.add_paragraph(line)
        
        i += 1
    
    # 保存到内存
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream



# 主界面
st.markdown('<h1 class="main-title">📚 教学设计智能体</h1>', unsafe_allow_html=True)
st.markdown('<p class="sub-title">为不同教学方法设计教学过程和PPT演示文稿</p>', unsafe_allow_html=True)

# 侧边栏 - API Key 配置
with st.sidebar:
    st.header("⚙️ 配置")
    api_key = st.text_input("DeepSeek API Key", type="password", help="输入你的 DeepSeek API Key")
    
    if api_key:
        st.session_state.client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com")
        st.success("✅ API Key 已配置")
    
    st.markdown("---")
    st.markdown("### 📖 使用说明")
    st.markdown("""
    1. 输入 API Key
    2. 填写完整的教学信息和设计要素
    3. 点击生成教学设计按钮
    4. 下载教学设计
    5. 点击生成PPT大纲按钮
    6. 下载PPT大纲
    7. 点击制作PPT演示前往Banana Slides
    """)
    
    st.markdown("---")
    st.markdown("### 💡 特色功能")
    st.markdown("""
    - 🎯 基于物理核心素养设计目标
    - 📝 自动生成板书设计
    - 👥 根据学生情况因材施教
    - ⏱️ 精确的时间分配
    - 📄 导出Word文档
    - 🎨 美观的界面设计
    """)

# 主内容区 - 单页面布局
st.markdown("## 📝 教学基本信息")
st.markdown("---")

col1, col2 = st.columns([1, 1])

with col1:
    subject = st.text_input("📚 学科/课程名称", placeholder="例如：高中物理")
    content = st.text_area("📖 授课内容", placeholder="例如：牛顿第二定律", height=100)
    audience = st.text_input("👥 教学对象", placeholder="例如：高一学生")
    
with col2:
    method = st.selectbox("🎓 教学方法", list(TEACHING_METHODS.keys()))
    st.info(f"💡 {TEACHING_METHODS[method]}")
    
    # 学生情况输入
    student_level = st.text_input(
        "🎓 学生情况",
        placeholder="例如：尖子班、普通班、基础薄弱班等",
        help="请描述学生的学习水平和特点"
    )

# 时间设置
st.markdown("#### ⏱️ 时间安排")
col_time1, col_time2 = st.columns([1, 1])

with col_time1:
    duration = st.slider("课时时长（分钟）", 30, 180, 45, step=5)
    st.caption(f"当前选择：{duration} 分钟")

with col_time2:
    custom_duration = st.number_input(
        "自定义时长（分钟）",
        min_value=10,
        max_value=300,
        value=None,
        step=5,
        placeholder="留空则使用滑块时长",
        help="如果需要精确时长，可以在这里输入"
    )
    if custom_duration:
        st.caption(f"将使用自定义时长：{custom_duration} 分钟")

# 教学设计核心要素
st.markdown("---")
st.markdown("## 🎯 教学设计核心要素")
st.markdown("---")

col3, col4 = st.columns([1, 1])

with col3:
    key_points = st.text_area(
        "🎯 教学重点",
        placeholder="请输入本节课的教学重点，例如：\n1. 理解牛顿第二定律的内容\n2. 掌握F=ma的应用",
        height=120
    )
    
    teaching_goals = st.text_area(
        "📋 期望的教学目标",
        placeholder="请输入您期望的教学目标（系统会结合物理核心素养进行优化）\n例如：\n1. 理解牛顿第二定律\n2. 能够应用定律解决问题",
        height=120,
        help="注意：最终输出会基于物理核心素养优化您的目标"
    )

with col4:
    difficulties = st.text_area(
        "🔍 教学难点",
        placeholder="请输入本节课的教学难点，例如：\n1. 加速度与力的关系理解\n2. 矢量性的把握",
        height=120
    )
    
    additional_requirements = st.text_area(
        "✨ 其他要求",
        placeholder="请输入其他教学要求，例如：\n- 使用多媒体演示\n- 安排小组讨论活动\n- 包含实验环节\n- 使用案例分析",
        height=120
    )

# 生成按钮
st.markdown("---")
col_btn1, col_btn2, col_btn3 = st.columns([1, 1, 1])

with col_btn1:
    generate_design = st.button("🎨 生成教学设计", use_container_width=True, type="primary", key="btn_design")

with col_btn2:
    generate_ppt = st.button("📊 生成PPT大纲", use_container_width=True, disabled=not st.session_state.get("teaching_design"), key="btn_ppt")

with col_btn3:
    st.markdown("""
    <div style='text-align: center; margin-top: 8px;'>
        <a href='https://bananaslides.online/' target='_blank' style='display: inline-block; padding: 0.5rem 2rem; background: linear-gradient(120deg, #F18F01, #FF6B35); color: white; text-decoration: none; border-radius: 8px; font-weight: 600;'>
            🎬 制作PPT演示
        </a>
    </div>
    """, unsafe_allow_html=True)


# 生成教学设计
if generate_design:
    if not api_key:
        st.error("❌ 请先在侧边栏输入 DeepSeek API Key")
    elif not subject or not content or not audience:
        st.error("❌ 请填写完整的基本教学信息（学科、内容、对象）")
    elif not student_level:
        st.error("❌ 请填写学生情况")
    elif not key_points or not difficulties:
        st.error("❌ 请填写教学重点和难点")
    elif not teaching_goals:
        st.error("❌ 请填写期望的教学目标")
    else:
        with st.spinner("🤔 正在生成教学设计..."):
            try:
                client = st.session_state.client
                
                # 立即设置标志，表示教学设计正在生成/已生成
                st.session_state.teaching_design = "generating"
                st.session_state.subject = subject
                st.session_state.content = content
                st.session_state.method = method
                
                response = generate_teaching_design(
                    client, subject, content, method, duration, audience,
                    key_points, difficulties, teaching_goals, student_level,
                    additional_requirements if additional_requirements else "无特殊要求",
                    custom_duration
                )
                
                st.markdown("---")
                st.markdown("## 📋 教学设计")
                design_container = st.empty()
                full_design = ""
                
                for chunk in response:
                    if chunk.choices[0].delta.content:
                        full_design += chunk.choices[0].delta.content
                        design_container.markdown(full_design)
                
                st.session_state.teaching_design = full_design
                st.session_state.subject = subject
                st.session_state.content = content
                st.session_state.method = method
                
                # 生成Word文档并保存到session_state
                word_file = markdown_to_word(full_design, f"{subject} - {content} 教学设计")
                st.session_state.teaching_design_word = word_file.read()
                st.session_state.design_subject = subject
                st.session_state.design_content = content
                
                st.success("✅ 教学设计生成完成！")
                
            except Exception as e:
                st.error(f"❌ 生成失败: {str(e)}")

# 教学设计下载按钮（在spinner外部，持久显示）
if st.session_state.get("teaching_design_word"):
    st.download_button(
        label="📥 下载教学设计（Word格式）",
        data=st.session_state.teaching_design_word,
        file_name=f"{st.session_state.get('design_subject', '教学')}_{st.session_state.get('design_content', '设计')}_教学设计.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
        key="download_design"
    )

# 生成PPT大纲
if generate_ppt:
    if not api_key:
        st.error("❌ 请先在侧边栏输入 DeepSeek API Key")
    else:
        with st.spinner("🤔 正在生成PPT大纲..."):
            try:
                client = st.session_state.client
                teaching_design = st.session_state.teaching_design
                subject = st.session_state.subject
                content = st.session_state.content
                method = st.session_state.method
                response = generate_ppt_outline(client, subject, content, method, teaching_design)
                
                st.markdown("---")
                st.markdown("## 📊 PPT演示文稿大纲")
                ppt_container = st.empty()
                full_ppt = ""
                
                for chunk in response:
                    if chunk.choices[0].delta.content:
                        full_ppt += chunk.choices[0].delta.content
                        ppt_container.markdown(full_ppt)
                
                st.session_state.ppt_outline = full_ppt
                st.success("✅ PPT大纲生成完成！")
                
                # 生成Word文档
                word_file = markdown_to_word(full_ppt, f"{subject} - {content} PPT大纲")
                
                # 下载按钮
                st.download_button(
                    label="📥 下载PPT大纲（Word格式）",
                    data=word_file,
                    file_name=f"{subject}_{content}_PPT大纲.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
                
                st.info("💡 提示：下载大纲后，可以点击上方'制作PPT演示'按钮，访问 Banana Slides 在线制作精美的演示文稿")
                
            except Exception as e:
                st.error(f"❌ 生成失败: {str(e)}")



# 显示历史结果
st.markdown("---")
if st.session_state.get("teaching_design") or st.session_state.get("ppt_outline"):
    st.markdown("## 📦 生成结果")

col_result1, col_result2 = st.columns([1, 1])

with col_result1:
    if st.session_state.get("teaching_design"):
        with st.expander("📋 查看教学设计", expanded=False):
            st.markdown(st.session_state.teaching_design)
            st.download_button(
                "💾 下载教学设计",
                st.session_state.teaching_design,
                file_name=f"{st.session_state.get('subject', '教学')}_{st.session_state.get('method', '')}教学设计.md",
                mime="text/markdown",
                use_container_width=True
            )

with col_result2:
    if st.session_state.get("ppt_outline"):
        with st.expander("📊 查看PPT大纲", expanded=False):
            st.markdown(st.session_state.ppt_outline)
            st.download_button(
                "💾 下载PPT大纲",
                st.session_state.ppt_outline,
                file_name=f"{st.session_state.get('subject', '教学')}_{st.session_state.get('method', '')}PPT大纲.md",
                mime="text/markdown",
                use_container_width=True
            )





# 页脚
st.markdown("---")
st.markdown("""
<div style='text-align: center; color: #666; padding: 20px;'>
    <p style='font-size: 1.1rem; margin-bottom: 0.5rem;'>💡 提示：生成的内容仅供参考，请根据实际情况调整</p>
    <p style='font-size: 0.9rem; color: #999;'>Powered by DeepSeek AI | 基于物理核心素养的教学设计 | 集成Banana Slides | 集成Banana Slides | 集成Banana Slides</p>
</div>
""", unsafe_allow_html=True)
